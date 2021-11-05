import get_vocab_list
import get_long_sentences
from docx import Document
from docx.oxml.ns import qn
import os

SOURCES_FOLDER_PATH = './sources'
RESULTS_FOLDER_PATH = './results'

PREVIEW_NUM = 20
PREVIEW_FILE_NAME = '阅读1-预习资料.docx'

QUIZ_NUM = 6
QUIZ_FILE_NAME = '阅读2-课首小测.docx'

SENTENCE_NUM = 2

TITLE_FONT = u'仿宋'
BODY_FONT = u'仿宋'

# Get all file names in the sources folder


def get_all_files(sources_folder_path):
    file_list = []
    for root, _, files in os.walk(sources_folder_path):
        for file in files:
            if file.endswith('.docx'):
                file_list.append(os.path.join(root, file))
    return file_list

# Use file names to read all every documents


def read_document(file_list):
    document_list = []
    for file in file_list:
        document_list.append(Document(file))
    return document_list

# Write preview file


def write_preview(select_vocab_list, results_folder_path, preview_file_name, title_font, body_font):
    preview = Document()

    run = preview.add_heading('', level=1).add_run(u'课前预习')
    run.font.name = title_font
    run._element.rPr.rFonts.set(qn('w:eastAsia'), title_font)

    preview.styles['Normal'].font.name = body_font
    preview.styles['Normal']._element.rPr.rFonts.set(
        qn('w:eastAsia'), body_font)

    for word in select_vocab_list:
        preview.add_paragraph(word)
    preview.save(os.path.join(results_folder_path, preview_file_name))

# Write quiz file


def write_quiz(quiz_vocab_list, longest_sentences, results_folder_path, quiz_file_name, title_font, body_font):
    quiz = Document()

    quiz.styles['Normal'].font.name = body_font
    quiz.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), body_font)

    run = quiz.add_heading('', level=1).add_run(u'课首小测')
    run.font.name = title_font
    run._element.rPr.rFonts.set(qn('w:eastAsia'), title_font)

    run = quiz.add_heading('', level=2).add_run(u'单词')
    run.font.name = title_font
    run._element.rPr.rFonts.set(qn('w:eastAsia'), title_font)
    for word in quiz_vocab_list:
        quiz.add_paragraph(word)

    run = quiz.add_heading('', level=2).add_run(u'长难句')
    run.font.name = title_font
    run._element.rPr.rFonts.set(qn('w:eastAsia'), title_font)
    for sentence in longest_sentences:
        quiz.add_paragraph(sentence)

    run = quiz.add_heading('', level=2).add_run(u'同义转换')
    run.font.name = title_font
    run._element.rPr.rFonts.set(qn('w:eastAsia'), title_font)

    quiz.save(os.path.join(results_folder_path, quiz_file_name))

# Read files


def read_files():
    file_list = get_all_files(SOURCES_FOLDER_PATH)
    document_list = read_document(file_list)
    return document_list


# Generate preview
def generate_preview(document_list):
    all_bold_vocab = get_vocab_list.get_all_bold_vocab(document_list)
    clean_vocab_list = get_vocab_list.clean_list(all_bold_vocab)
    select_vocab_list = get_vocab_list.select_words(
        clean_vocab_list, PREVIEW_NUM)
    write_preview(select_vocab_list, RESULTS_FOLDER_PATH,
                  PREVIEW_FILE_NAME, TITLE_FONT, BODY_FONT)
    return select_vocab_list


# Generate quiz
def generate_quiz(document_list, select_vocab_list):
    quiz_vocab_list = get_vocab_list.select_words(select_vocab_list, QUIZ_NUM)
    all_text = get_long_sentences.get_all_text(document_list)
    all_sentences = get_long_sentences.split_text_to_sentences(all_text)
    longest_sentences = get_long_sentences.select_longest_sentences(
        all_sentences, SENTENCE_NUM)
    write_quiz(quiz_vocab_list, longest_sentences,
               RESULTS_FOLDER_PATH, QUIZ_FILE_NAME, TITLE_FONT, BODY_FONT)


# Main
document_list = read_files()
select_vocab_list = generate_preview(document_list)
generate_quiz(document_list, select_vocab_list)
