from docx import Document

FILE_PATH_1 = './test1.docx'
FILE_PATH_2 = './test2.docx'
SENTENCE_NUM = 5
RESULT_FILE_NAME = 'reading.docx'


def get_all_text(file_path_1, file_path_2):
    all_text = []

    document_1 = Document(file_path_1)
    for para in document_1.paragraphs:
        for run in para.runs:
            all_text.append(run.text)

    document_2 = Document(file_path_2)
    for para in document_2.paragraphs:
        for run in para.runs:
            all_text.append(run.text)
    return all_text


def split_text_to_sentences(all_text):
    all_sentences = []
    # get rid of the title
    for para in all_text[1:]:
        sentences = para.split('.')
        for sentence in sentences:
            if sentence != '':
                sentence += '.'
                all_sentences.append(sentence)
    return all_sentences


def select_longest_sentences(all_sentences, sentence_num):
    sort_list = sorted(all_sentences, key=len, reverse=True)
    return sort_list[:sentence_num]


def write_sentences_to_file(longest_sentences, file_name):
    result = Document()
    for sentence in longest_sentences:
        result.add_paragraph(sentence)
    result.save(file_name)


all_text = get_all_text(FILE_PATH_1, FILE_PATH_2)
all_sentences = split_text_to_sentences(all_text)
longest_sentences = select_longest_sentences(all_sentences, SENTENCE_NUM)
write_sentences_to_file(longest_sentences, RESULT_FILE_NAME)
